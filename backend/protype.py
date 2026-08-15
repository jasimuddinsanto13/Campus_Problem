import docx

doc = docx.Document()

# Title
doc.add_heading('NITER-Pulse: Smart Classroom Discovery & Dynamic Room Booking System', level=0)
doc.add_paragraph('Software Requirements Specification (SRS) & Prototype Design Document\n')

# Section 1
doc.add_heading('1. Executive Summary', level=1)
doc.add_heading('1.1 Project Title', level=2)
doc.add_paragraph('NITER-Pulse (Smart Classroom Discovery & Dynamic Room Booking Module)')

doc.add_heading('1.2 Problem Statement', level=2)
doc.add_paragraph(
    "Campus operations at NITER experience severe classroom shortages during peak academic hours. "
    "Finding an empty classroom for rescheduled lectures, extra classes, or examinations requires "
    "manual physical checks or tedious coordination across departments.\n\n"
    "Furthermore, when external batches or departments occupy classrooms for examinations, "
    "other faculty members struggle to locate available spaces for their regular or makeup sessions."
)

doc.add_heading('1.3 Proposed Solution', level=2)
doc.add_paragraph(
    "NITER-Pulse provides a real-time, centralized platform where:\n"
    "1. Public Free Room Visibility: Both students and teachers can view all currently empty classrooms along with the exact duration they remain unassigned (e.g., 'Room 302: Free from 10:30 AM to 01:30 PM').\n"
    "2. Student-to-Teacher Loop: Students can check available slots and inform their course teachers about potential rooms for extra or rescheduled classes.\n"
    "3. Faculty Booking Privilege: Teachers have exclusive authorization to select an available room, check its free time window, and instantly book it for makeup classes, extra sessions, or examinations.\n"
    "4. Conflict Prevention: Upon booking, the room is locked in the central system, automatically removing it from the available pool and notifying enrolled students."
)

# Section 2
doc.add_heading('2. System Architecture & Workflow', level=1)
doc.add_heading('2.1 User Roles', level=2)
doc.add_paragraph(
    "- Student (View-Only Access): Can view real-time free room lists, room capacity, and 'free until' time slots.\n"
    "- Faculty / Teacher (Admin & Booking Access): Can view free rooms, select target time windows, and execute instant room bookings for extra classes, reschedules, or exams.\n"
    "- System Administrator: Oversees master schedules, manages room capacities, and overrides invalid bookings."
)

doc.add_heading('2.2 End-to-End Workflow Scenario', level=2)
doc.add_paragraph(
    "1. Room Discovery: A teacher needs to schedule a 1.5-hour exam for CSE 13th Batch. Alternatively, students notice Room 204 is empty for 2 hours and suggest it to their teacher.\n"
    "2. Search & Verification: The teacher opens the NITER-Pulse dashboard, inputs the required date and time window (e.g., Today, 11:30 AM – 01:00 PM), and reviews all rooms marked 'Free'.\n"
    "3. Reservation: The teacher selects Room 204, chooses the booking type ('Examination / Reschedule / Extra Class'), and clicks 'Confirm Booking'.\n"
    "4. State Update & Notification: The database locks Room 204 for that period, updates the public dashboard so no other department can double-book it, and dispatches a notification to the target batch."
)

# Section 3
doc.add_heading('3. Database Schema (Django ORM Models)', level=1)
code_models = """from django.db import models
from django.contrib.auth.models import AbstractUser

class User(AbstractUser):
    ROLE_CHOICES = (
        ('student', 'Student'),
        ('teacher', 'Teacher'),
        ('admin', 'Admin'),
    )
    role = models.CharField(max_length=20, choices=ROLE_CHOICES, default='student')
    department = models.CharField(max_length=50, blank=True)
    section = models.CharField(max_length=10, blank=True)
    batch = models.CharField(max_length=20, blank=True)
    phone_number = models.CharField(max_length=15, blank=True)

class Room(models.Model):
    room_number = models.CharField(max_length=20, unique=True)
    building = models.CharField(max_length=100, default="Academic Building")
    capacity = models.IntegerField(default=60)

class Routine(models.Model):
    DAYS = (
        ('MON', 'Monday'), ('TUE', 'Tuesday'), ('WED', 'Wednesday'),
        ('THU', 'Thursday'), ('FRI', 'Friday'), ('SAT', 'Saturday'), ('SUN', 'Sunday'),
    )
    teacher = models.ForeignKey(User, on_delete=models.CASCADE)
    subject = models.CharField(max_length=100)
    department = models.CharField(max_length=50)
    section = models.CharField(max_length=10)
    room = models.ForeignKey(Room, on_delete=models.CASCADE)
    day = models.CharField(max_length=3, choices=DAYS)
    start_time = models.TimeField()
    end_time = models.TimeField()

class RoomBooking(models.Model):
    BOOKING_TYPES = (
        ('extra_class', 'Extra Class'),
        ('reschedule', 'Rescheduled Class'),
        ('exam', 'Examination'),
    )
    room = models.ForeignKey(Room, on_delete=models.CASCADE)
    booked_by = models.ForeignKey(User, on_delete=models.CASCADE)
    booking_type = models.CharField(max_length=20, choices=BOOKING_TYPES)
    department = models.CharField(max_length=50)
    batch_section = models.CharField(max_length=50)
    date = models.DateField()
    start_time = models.TimeField()
    end_time = models.TimeField()
    created_at = models.DateTimeField(auto_now_add=True)"""

doc.add_paragraph(code_models)

# Save
doc.save("NITER_Pulse_Specification.docx")
print("SUCCESS: File generated as 'NITER_Pulse_Specification.docx'")