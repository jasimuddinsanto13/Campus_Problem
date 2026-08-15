import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_niter_pulse_srs():
    doc = Document()

    # Document Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    run_title = title_p.add_run("NITER-Pulse: Smart Classroom Discovery & Dynamic Room Booking System")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 32, 67)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_p.add_run("Software Requirements Specification (SRS) & Prototype Design Document")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(90, 105, 120)

    # Helper function for headings
    def add_sec_heading(text, level=1):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run(text)
        run.font.name = "Arial"
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(15, 32, 67)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(30, 60, 110)
        return h

    # Section 1
    add_sec_heading("1. Executive Summary", level=1)
    
    add_sec_heading("1.1 Project Title", level=2)
    doc.add_paragraph("NITER-Pulse (Smart Classroom Discovery & Dynamic Room Booking Module)")

    add_sec_heading("1.2 Problem Statement", level=2)
    doc.add_paragraph(
        "Campus operations at NITER experience severe classroom shortages during peak academic hours. "
        "Finding an empty classroom for rescheduled lectures, extra classes, or examinations requires manual physical "
        "checks or tedious coordination across departments. Furthermore, when external batches or departments occupy "
        "classrooms for examinations, other faculty members struggle to locate available spaces for their regular or makeup sessions."
    )

    add_sec_heading("1.3 Proposed Solution", level=2)
    p_sol = doc.add_paragraph()
    solutions = [
        "Public Free Room Visibility: Both students and teachers can view all currently empty classrooms along with the exact duration they remain unassigned (e.g., 'Room 302: Free from 10:30 AM to 01:30 PM').",
        "Student-to-Teacher Loop: Students can check available slots and inform their course teachers about potential rooms for extra or rescheduled classes.",
        "Faculty Booking Privilege: Teachers have exclusive authorization to select an available room, check its free time window, and instantly book it for makeup classes, extra sessions, or examinations.",
        "Conflict Prevention: Upon booking, the room is locked in the central system, automatically removing it from the available pool and notifying enrolled students.",
        "Administrative Controls & Routine Management: System administrators manage master academic routines, approve pending registration requests, and view detailed user profiles for students and teachers."
    ]
    for idx, item in enumerate(solutions, 1):
        doc.add_paragraph(f"{idx}. {item}", style='List Number')

    # Section 2
    add_sec_heading("2. System Architecture & Workflow", level=1)

    add_sec_heading("2.1 User Roles", level=2)
    roles = [
        ("Student (View-Only Access)", "Can view real-time free room lists, room capacity, and 'free until' time slots. Requires admin approval upon account registration."),
        ("Faculty / Teacher (Admin & Booking Access)", "Can view free rooms, select target time windows, and execute instant room bookings for extra classes, reschedules, or exams. Requires admin approval upon account registration."),
        ("System Administrator", "Oversees master schedules, uploads department routines, manages room capacities, overrides invalid bookings, approves pending registration requests, and manages user accounts.")
    ]
    for role_title, role_desc in roles:
        p = doc.add_paragraph(style='List Bullet')
        r_bold = p.add_run(f"{role_title}: ")
        r_bold.bold = True
        p.add_run(role_desc)

    add_sec_heading("2.2 End-to-End Workflow Scenario", level=2)
    
    doc.add_paragraph().add_run("A. Registration & Approval Workflow").bold = True
    reg_steps = [
        "Sign-Up Request: A new user (Student/Teacher) creates an account. The account remains in 'pending' status by default.",
        "Admin Review: The System Administrator logs into the Admin Dashboard and opens the Pending Registrations section.",
        "Approval Decision: The admin clicks 'Approve' or 'Reject'. Approved users gain immediate system access."
    ]
    for idx, s in enumerate(reg_steps, 1):
        doc.add_paragraph(f"{idx}. {s}", style='List Number')

    doc.add_paragraph().add_run("B. Master Routine Setup Workflow").bold = True
    rt_steps = [
        "Parameter Selection: Admin selects the Department, Batch, and Subject from the top selection options.",
        "Timetable Filling: Admin inputs schedule details across the Sunday to Thursday table grid.",
        "Activation: The saved routine establishes baseline occupancy for room availability searches."
    ]
    for idx, s in enumerate(rt_steps, 1):
        doc.add_paragraph(f"{idx}. {s}", style='List Number')

    doc.add_paragraph().add_run("C. Room Discovery & Booking Workflow").bold = True
    bk_steps = [
        "Room Discovery: A teacher needs to schedule a 1.5-hour exam or makeup class.",
        "Search & Verification: The teacher queries the dashboard for a specific target time window and checks empty rooms.",
        "Reservation: The teacher selects an available room, defines the booking reason, and confirms the reservation.",
        "State Lock & Notification: The room is locked in real-time and notification is dispatched to students."
    ]
    for idx, s in enumerate(bk_steps, 1):
        doc.add_paragraph(f"{idx}. {s}", style='List Number')

    # Section 3
    add_sec_heading("3. Administrative Modules & Dashboard Requirements", level=1)

    add_sec_heading("3.1 Registration Approval Queue", level=2)
    doc.add_paragraph(
        "A dedicated section on the Admin Dashboard lists all pending user registrations. "
        "The admin can view applicant details (Name, Role, Email, Department, Phone) and execute Approve or Reject actions."
    )

    add_sec_heading("3.2 User Info Section", level=2)
    doc.add_paragraph(
        "Provides a two-tab interface listing registered user information:"
    )
    u_info = [
        "Student View: Displays student Username, Email Address, Password (hashed/represented), Department, Batch, Section, and Phone Number.",
        "Teacher View: Displays teacher Username, Email Address, Password (hashed/represented), Department, and Phone Number."
    ]
    for item in u_info:
        doc.add_paragraph(item, style='List Bullet')

    add_sec_heading("3.3 Routine Setup & Entry Module", level=2)
    doc.add_paragraph("Includes top parameter dropdowns and a structured schedule grid:")
    
    # Top selections description
    p_top = doc.add_paragraph(style='List Bullet')
    p_top.add_run("Header Parameters: ").bold = True
    p_top.add_run("Department Name | Batch Name | Subject Name")

    # Table representation for Routine
    doc.add_paragraph().add_run("Sample Weekly Routine Entry Table Format (Sun - Thu):").italic = True
    
    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ["Day", "Time Slot", "Course Name", "Teacher Name"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1F497D")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    sample_data = [
        ("Sunday", "09:00 AM - 10:30 AM", "CSE-101 (Data Structures)", "Dr. Ahsan Habib"),
        ("Monday", "10:30 AM - 12:00 PM", "CSE-103 (Algorithms)", "Prof. Selim Reza"),
        ("Tuesday", "01:30 PM - 03:00 PM", "CSE-105 (Database Systems)", "Tania Sultana"),
        ("Wednesday", "09:00 AM - 10:30 AM", "CSE-107 (Software Eng.)", "Dr. Ahsan Habib"),
        ("Thursday", "11:30 AM - 01:00 PM", "CSE-109 (Web Tech)", "Mahmudul Hasan"),
    ]

    for row_idx, data in enumerate(sample_data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F2F2F2")

    # Section 4
    add_sec_heading("4. Database Schema (Django ORM Models)", level=1)

    code_snippet = (
        "from django.db import models\n"
        "from django.contrib.auth.models import AbstractUser\n\n"
        "class User(AbstractUser):\n"
        "    ROLE_CHOICES = (\n"
        "        ('student', 'Student'),\n"
        "        ('teacher', 'Teacher'),\n"
        "        ('admin', 'Admin'),\n"
        "    )\n"
        "    STATUS_CHOICES = (\n"
        "        ('pending', 'Pending Approval'),\n"
        "        ('approved', 'Approved'),\n"
        "        ('rejected', 'Rejected'),\n"
        "    )\n"
        "    role = models.CharField(max_length=20, choices=ROLE_CHOICES, default='student')\n"
        "    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='pending')\n"
        "    department = models.CharField(max_length=50, blank=True)\n"
        "    section = models.CharField(max_length=10, blank=True)\n"
        "    batch = models.CharField(max_length=20, blank=True)\n"
        "    phone_number = models.CharField(max_length=15, blank=True)\n\n"
        "class Room(models.Model):\n"
        "    room_number = models.CharField(max_length=20, unique=True)\n"
        "    building = models.CharField(max_length=100, default=\"Academic Building\")\n"
        "    capacity = models.IntegerField(default=60)\n\n"
        "class Routine(models.Model):\n"
        "    DAYS_OF_WEEK = (\n"
        "        ('SUN', 'Sunday'), ('MON', 'Monday'), ('TUE', 'Tuesday'),\n"
        "        ('WED', 'Wednesday'), ('THU', 'Thursday'),\n"
        "    )\n"
        "    department = models.CharField(max_length=50)\n"
        "    batch = models.CharField(max_length=20)\n"
        "    subject = models.CharField(max_length=100)\n"
        "    teacher = models.ForeignKey(User, on_delete=models.CASCADE, limit_choices_to={'role': 'teacher'})\n"
        "    room = models.ForeignKey(Room, on_delete=models.CASCADE)\n"
        "    day = models.CharField(max_length=3, choices=DAYS_OF_WEEK)\n"
        "    start_time = models.TimeField()\n"
        "    end_time = models.TimeField()\n\n"
        "class RoomBooking(models.Model):\n"
        "    BOOKING_TYPES = (\n"
        "        ('extra_class', 'Extra Class'),\n"
        "        ('reschedule', 'Rescheduled Class'),\n"
        "        ('exam', 'Examination'),\n"
        "    )\n"
        "    room = models.ForeignKey(Room, on_delete=models.CASCADE)\n"
        "    booked_by = models.ForeignKey(User, on_delete=models.CASCADE)\n"
        "    booking_type = models.CharField(max_length=20, choices=BOOKING_TYPES)\n"
        "    department = models.CharField(max_length=50)\n"
        "    batch_section = models.CharField(max_length=50)\n"
        "    date = models.DateField()\n"
        "    start_time = models.TimeField()\n"
        "    end_time = models.TimeField()\n"
        "    created_at = models.DateTimeField(auto_now_add=True)\n"
    )

    code_p = doc.add_paragraph()
    code_run = code_p.add_run(code_snippet)
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(9.5)
    code_run.font.color.rgb = RGBColor(40, 40, 40)

    # Save to workspace root
    output_filename = "NITER_Pulse_Specification.docx"
    doc.save(output_filename)
    print(f"Successfully generated/updated document: {os.path.abspath(output_filename)}")

if __name__ == "__main__":
    create_niter_pulse_srs()